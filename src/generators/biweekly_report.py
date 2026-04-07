"""
半月报生成模块 - 汇总标记的重要新闻
"""
import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from ..fetchers.base import NewsArticle


class BiweeklyReportGenerator:
    """半月报生成器"""

    def __init__(self, config: dict, logger=None):
        self.config = config
        self.logger = logger
        output_config = config.get("output", {})
        self.biweekly_dir = Path(output_config.get("biweekly_dir", "output/biweekly"))
        self.data_dir = Path(output_config.get("data_dir", "data"))
        self.date_format = output_config.get("date_format", "%Y%m%d")
        self.marked_file = self.data_dir / "marked_news.json"

    def mark_news(self, article: NewsArticle, analysis: dict = None):
        """标记一条新闻为重要（添加到半月报候选）"""
        marked = self._load_marked()

        # 检查是否已标记
        for item in marked:
            if item["url"] == article.url:
                if self.logger:
                    self.logger.info(f"[半月报] 新闻已标记: {article.title[:30]}")
                return

        entry = article.to_dict()
        if analysis:
            entry["analysis"] = analysis
        entry["marked_at"] = datetime.now().isoformat()

        marked.append(entry)
        self._save_marked(marked)

        if self.logger:
            self.logger.info(f"[半月报] 已标记重要新闻: {article.title[:30]}")

    def unmark_news(self, url: str):
        """取消标记"""
        marked = self._load_marked()
        marked = [item for item in marked if item["url"] != url]
        self._save_marked(marked)

    def get_marked_news(self, start_date: datetime = None,
                        end_date: datetime = None) -> List[dict]:
        """获取标记的新闻（可按日期范围过滤）"""
        marked = self._load_marked()

        if start_date:
            marked = [
                item for item in marked
                if item.get("publish_time") and datetime.fromisoformat(item["publish_time"]) >= start_date
            ]
        if end_date:
            marked = [
                item for item in marked
                if item.get("publish_time") and datetime.fromisoformat(item["publish_time"]) <= end_date
            ]

        return marked

    def generate(self, start_date: datetime = None,
                 end_date: datetime = None) -> Dict[str, str]:
        """
        生成半月报

        Args:
            start_date: 半月报起始日期
            end_date: 半月报结束日期

        Returns:
            {"docx": "文件路径", "md": "文件路径"}
        """
        if start_date is None:
            # 默认取本月1号到15号，或16号到月末
            today = datetime.now()
            if today.day <= 15:
                start_date = datetime(today.year, today.month, 1)
                end_date = datetime(today.year, today.month, 15)
            else:
                start_date = datetime(today.year, today.month, 16)
                if today.month == 12:
                    end_date = datetime(today.year + 1, 1, 1) - timedelta(days=1)
                else:
                    end_date = datetime(today.year, today.month + 1, 1) - timedelta(days=1)

        # 获取标记的新闻
        marked = self.get_marked_news(start_date, end_date)

        if not marked:
            if self.logger:
                self.logger.warning(f"[半月报] {start_date.strftime('%m.%d')}-{end_date.strftime('%m.%d')} 无标记新闻")
            return {}

        # 确保输出目录存在
        self.biweekly_dir.mkdir(parents=True, exist_ok=True)

        start_str = start_date.strftime(self.date_format)
        end_str = end_date.strftime(self.date_format)
        display_period = f"{start_date.strftime('%Y年%m月%d日')} - {end_date.strftime('%m月%d日')}"

        # 生成Word文档
        docx_path = self.biweekly_dir / f"洞察信息半月报_{start_str}_{end_str}.docx"
        self._generate_docx(docx_path, marked, display_period)

        # 生成Markdown文件
        md_path = self.biweekly_dir / f"洞察信息半月报_{start_str}_{end_str}.md"
        self._generate_markdown(md_path, marked, display_period)

        return {
            "docx": str(docx_path),
            "md": str(md_path),
        }

    def _generate_docx(self, filepath: Path, marked: List[dict], display_period: str):
        """生成半月报Word文档"""
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(10)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 标题
        title = doc.add_heading("", level=0)
        run = title.add_run("洞察信息半月报")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 日期范围
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"汇总周期：{display_period}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)

        doc.add_paragraph()

        # 按层面分类分组
        grouped = {}
        for item in marked:
            cat = item.get("category", "行业/市场层")
            if cat not in grouped:
                grouped[cat] = []
            grouped[cat].append(item)

        category_order = ["国家/政策层", "行业/市场层", "技术/研发层", "业务/竞争层"]

        for category in category_order:
            items = grouped.get(category, [])
            if not items:
                continue

            doc.add_heading(category, level=1)

            table = doc.add_table(rows=1, cols=6)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ["日期", "事件内容", "参与方", "事件洞察", "对岚图的影响及启示", "信息来源"]
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.name = "微软雅黑"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            for item in items:
                row = table.add_row()
                analysis = item.get("analysis", {})
                pub_time = item.get("publish_time", "")
                if pub_time:
                    try:
                        dt = datetime.fromisoformat(pub_time)
                        date_str = dt.strftime("%m.%d")
                    except ValueError:
                        date_str = ""
                else:
                    date_str = ""

                values = [
                    date_str,
                    analysis.get("事件内容", item.get("title", "")[:60]),
                    analysis.get("参与方", ""),
                    analysis.get("事件洞察", ""),
                    analysis.get("对岚图的影响及启示", ""),
                    item.get("url", ""),
                ]
                for i, value in enumerate(values):
                    cell = row.cells[i]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            doc.add_paragraph()

        doc.save(str(filepath))
        if self.logger:
            self.logger.info(f"[半月报] Word文档已生成: {filepath}")

    def _generate_markdown(self, filepath: Path, marked: List[dict], display_period: str):
        """生成半月报Markdown文件"""
        lines = []
        lines.append(f"# 洞察信息半月报")
        lines.append(f"\n**汇总周期**：{display_period}\n")

        grouped = {}
        for item in marked:
            cat = item.get("category", "行业/市场层")
            if cat not in grouped:
                grouped[cat] = []
            grouped[cat].append(item)

        category_order = ["国家/政策层", "行业/市场层", "技术/研发层", "业务/竞争层"]

        for category in category_order:
            items = grouped.get(category, [])
            if not items:
                continue

            lines.append(f"\n## {category}\n")
            lines.append("| 日期 | 事件内容 | 参与方 | 事件洞察 | 对岚图的影响及启示 | 信息来源 |")
            lines.append("|------|----------|--------|----------|-------------------|----------|")

            for item in items:
                analysis = item.get("analysis", {})
                pub_time = item.get("publish_time", "")
                if pub_time:
                    try:
                        dt = datetime.fromisoformat(pub_time)
                        date_str = dt.strftime("%m.%d")
                    except ValueError:
                        date_str = ""
                else:
                    date_str = ""

                event = analysis.get("事件内容", item.get("title", "")[:60]).replace("|", "\\|")
                parties = analysis.get("参与方", "").replace("|", "\\|")
                insight = analysis.get("事件洞察", "").replace("|", "\\|")
                suggestion = analysis.get("对岚图的影响及启示", "").replace("|", "\\|")
                url = item.get("url", "")

                lines.append(
                    f"| {date_str} | {event} | {parties} | {insight} "
                    f"| {suggestion} | [链接]({url}) |"
                )

        content = "\n".join(lines)
        filepath.write_text(content, encoding="utf-8")
        if self.logger:
            self.logger.info(f"[半月报] Markdown文件已生成: {filepath}")

    def _load_marked(self) -> List[dict]:
        """加载标记的新闻"""
        if not self.marked_file.exists():
            return []
        try:
            with open(self.marked_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return []

    def _save_marked(self, marked: List[dict]):
        """保存标记的新闻"""
        self.marked_file.parent.mkdir(parents=True, exist_ok=True)
        with open(self.marked_file, "w", encoding="utf-8") as f:
            json.dump(marked, f, ensure_ascii=False, indent=2)

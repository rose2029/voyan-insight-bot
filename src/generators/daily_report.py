"""
日报生成模块 - Word(.docx) 和 Markdown 输出
"""
import json
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from ..fetchers.base import NewsArticle


# 层面分类顺序及子标题
CATEGORY_CONFIG = {
    "国家/政策层": "国家政策导向",
    "行业/市场层": "行业市场动态",
    "技术/研发层": "电池、智能驾驶、智能座舱、芯片、新材料等领域的技术突破",
    "业务/竞争层": "车企动作",
}


class DailyReportGenerator:
    """日报生成器"""

    def __init__(self, config: dict, logger=None):
        self.config = config
        self.logger = logger
        output_config = config.get("output", {})
        self.daily_dir = Path(output_config.get("daily_dir", "output/daily"))
        self.archive_dir = Path(output_config.get("archive_dir", "output/archive"))
        self.date_format = output_config.get("date_format", "%Y%m%d")

    def generate(self, articles: List[NewsArticle], analysis_results: List[dict],
                 summary: dict, report_date: datetime = None) -> Dict[str, str]:
        """
        生成日报（Word + Markdown）

        Args:
            articles: 筛选后的新闻列表
            analysis_results: 每条新闻的LLM分析结果
            summary: 日报总结（要点概括、战略意义等）
            report_date: 日报日期

        Returns:
            {"docx": "文件路径", "md": "文件路径"}
        """
        if report_date is None:
            report_date = datetime.now()

        date_str = report_date.strftime(self.date_format)
        display_date = report_date.strftime("%Y年%m月%d日")

        # 按层面分类分组
        grouped = self._group_by_category(articles, analysis_results)

        # 确保输出目录存在
        self.daily_dir.mkdir(parents=True, exist_ok=True)

        # 生成Word文档
        docx_path = self.daily_dir / f"洞察信息收集日报_{date_str}.docx"
        self._generate_docx(docx_path, grouped, summary, display_date, date_str)

        # 生成Markdown文件
        md_path = self.daily_dir / f"洞察信息收集日报_{date_str}.md"
        self._generate_markdown(md_path, grouped, summary, display_date, date_str)

        return {
            "docx": str(docx_path),
            "md": str(md_path),
        }

    def _group_by_category(self, articles: List[NewsArticle],
                           analysis_results: List[dict]) -> Dict[str, List[tuple]]:
        """按层面分类分组"""
        category_order = list(CATEGORY_CONFIG.keys())
        grouped = {cat: [] for cat in category_order}
        grouped.setdefault("其他", [])

        for article, analysis in zip(articles, analysis_results):
            cat = article.category or "行业/市场层"
            if cat not in grouped:
                cat = "其他"
            grouped[cat].append((article, analysis))

        # 移除空分类
        return {k: v for k, v in grouped.items() if v}

    def _generate_docx(self, filepath: Path, grouped: dict, summary: dict,
                       display_date: str, date_str: str):
        """生成Word文档"""
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(10)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 标题
        title = doc.add_heading("", level=0)
        run = title.add_run(f"洞察信息收集日报")
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(display_date)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(102, 102, 102)

        doc.add_paragraph()  # 空行

        # 今日要点概括（放在最前面）
        if summary and summary.get("要点概括"):
            doc.add_heading("今日要点概括", level=1)
            for i, item in enumerate(summary["要点概括"], 1):
                para = doc.add_paragraph()
                run = para.add_run(f"{i}. {item}")
                run.font.size = Pt(10)
                run.font.name = "微软雅黑"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            doc.add_paragraph()  # 空行

        # 各层面分类新闻表格
        for category, items in grouped.items():
            doc.add_heading(category, level=1)

            if not items:
                doc.add_paragraph("暂无相关新闻")
                continue

            # 子标题行（参考模板格式）
            sub_title = CATEGORY_CONFIG.get(category, "")
            if sub_title:
                sub_para = doc.add_paragraph()
                run = sub_para.add_run(sub_title)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.font.italic = True

            # 创建表格
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 表头
            headers = ["时间", "事件内容", "参与方", "事件影响", "事件洞察",
                       "对岚图的影响及启示", "信息来源"]
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(9)
                        run.font.name = "微软雅黑"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 数据行
            for article, analysis in items:
                row = table.add_row()
                time_str = article.publish_date_str or date_str[4:6] + "." + date_str[6:8]
                values = [
                    time_str,
                    analysis.get("事件内容", article.title),
                    analysis.get("参与方", ""),
                    analysis.get("事件影响", ""),
                    analysis.get("事件洞察", ""),
                    analysis.get("对岚图的影响及启示", ""),
                    article.url,
                ]
                for i, value in enumerate(values):
                    cell = row.cells[i]
                    cell.text = str(value)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = "微软雅黑"
                            run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            # 设置列宽（参考模板比例）
            widths = [Cm(1.2), Cm(4.0), Cm(1.8), Cm(3.5), Cm(3.5), Cm(3.0), Cm(2.5)]
            for row in table.rows:
                for i, width in enumerate(widths):
                    row.cells[i].width = width

            doc.add_paragraph()  # 空行

        # 总结模块（战略意义、风险预警、近期关注）
        if summary:
            self._add_summary_section(doc, summary)

        doc.save(str(filepath))
        if self.logger:
            self.logger.info(f"[日报] Word文档已生成: {filepath}")

    def _add_summary_section(self, doc: Document, summary: dict):
        """添加总结模块到Word文档（参考模板格式）"""
        doc.add_page_break()
        doc.add_heading("初步分析与预警", level=1)

        sections = [
            ("战略意义", summary.get("战略意义", [])),
            ("风险预警", summary.get("风险预警", [])),
            ("近期关注", summary.get("近期关注", [])),
        ]

        for section_title, items in sections:
            doc.add_heading(section_title, level=2)
            if items:
                for i, item in enumerate(items, 1):
                    para = doc.add_paragraph()
                    run = para.add_run(f"{item}")
                    run.font.size = Pt(10)
                    run.font.name = "微软雅黑"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            else:
                doc.add_paragraph("（暂无）")

    def _generate_markdown(self, filepath: Path, grouped: dict, summary: dict,
                           display_date: str, date_str: str):
        """生成Markdown文件（参考模板格式）"""
        lines = []
        lines.append(f"# 洞察信息收集日报")
        lines.append(f"\n**日期**：{display_date}\n")

        # 今日要点概括（放在最前面）
        if summary and summary.get("要点概括"):
            lines.append("## 今日要点概括\n")
            for i, item in enumerate(summary["要点概括"], 1):
                lines.append(f"{i}. {item}")
            lines.append("")

        # 各层面分类新闻
        for category, items in grouped.items():
            lines.append(f"\n## {category}\n")

            if not items:
                lines.append("暂无相关新闻\n")
                continue

            # 子标题
            sub_title = CATEGORY_CONFIG.get(category, "")
            if sub_title:
                lines.append(f"**{sub_title}**\n")

            # 表头
            lines.append("| 时间 | 事件内容 | 参与方 | 事件影响 | 事件洞察 | 对岚图的影响及启示 | 信息来源 |")
            lines.append("|------|----------|--------|----------|----------|-------------------|----------|")

            for article, analysis in items:
                time_str = article.publish_date_str or date_str[4:6] + "." + date_str[6:8]
                event = analysis.get("事件内容", article.title).replace("|", "\\|")
                parties = analysis.get("参与方", "").replace("|", "\\|")
                impact = analysis.get("事件影响", "").replace("|", "\\|")
                insight = analysis.get("事件洞察", "").replace("|", "\\|")
                suggestion = analysis.get("对岚图的影响及启示", "").replace("|", "\\|")
                url = article.url

                lines.append(
                    f"| {time_str} | {event} | {parties} | {impact} "
                    f"| {insight} | {suggestion} | [链接]({url}) |"
                )

        # 总结模块（战略意义、风险预警、近期关注）
        if summary:
            lines.append("\n---\n")
            lines.append("## 初步分析与预警\n")

            sections = [
                ("战略意义", summary.get("战略意义", [])),
                ("风险预警", summary.get("风险预警", [])),
                ("近期关注", summary.get("近期关注", [])),
            ]

            for section_title, items in sections:
                lines.append(f"### {section_title}\n")
                if items:
                    for i, item in enumerate(items, 1):
                        lines.append(f"{i}. {item}")
                else:
                    lines.append("（暂无）")
                lines.append("")

        content = "\n".join(lines)
        filepath.write_text(content, encoding="utf-8")
        if self.logger:
            self.logger.info(f"[日报] Markdown文件已生成: {filepath}")
